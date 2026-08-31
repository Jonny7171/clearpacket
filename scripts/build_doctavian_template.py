#!/usr/bin/env python3
"""Build the Doctavian supplier-credit template from the official sample."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "111111"
ACCENT = "4743FF"
MUTED = "666666"
PALE = "F0EFFF"


def set_run_font(run, name: str = "Poppins", size: float = 10.5, *, bold: bool = False, color: str = INK):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.12):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_text(doc, text: str, *, size: float = 10.5, bold: bool = False, color: str = INK,
             before: float = 0, after: float = 6, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after)
    set_run_font(paragraph.add_run(text), size=size, bold=bold, color=color)
    return paragraph


def add_section_label(doc, text: str):
    paragraph = add_text(doc, text.upper(), size=9.5, bold=True, before=9, after=3)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def shade_paragraph(paragraph, fill: str, *, border: str | None = None):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    if border:
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), border)
        borders.append(bottom)
        properties.append(borders)


def add_expression(doc, text: str, *, accent: bool = False, after: float = 4):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.05)
    set_run_font(paragraph.add_run(text), size=9.2, color=ACCENT if accent else INK)
    return paragraph


def reset_body_to_first_section(doc: Document):
    body = doc._element.body
    section_properties = body.xpath(".//w:sectPr")
    if not section_properties:
        raise RuntimeError("The reference template has no section properties.")
    first_section = deepcopy(section_properties[0])
    for child in list(body):
        body.remove(child)
    body.append(first_section)


def build(reference: Path, output: Path):
    doc = Document(reference)
    reset_body_to_first_section(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.38)
    columns = section._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section._sectPr.append(columns)
    for column in list(columns):
        columns.remove(column)
    columns.set(qn("w:num"), "1")
    columns.set(qn("w:equalWidth"), "1")
    columns.attrib.pop(qn("w:sep"), None)

    identity = doc.add_paragraph()
    set_paragraph_spacing(identity, after=1)
    identity.paragraph_format.tab_stops.add_tab_stop(Inches(6.75))
    set_run_font(identity.add_run("Northwind Fabrication"), name="Poppins SemiBold", size=12.5, bold=True)
    identity.add_run("\t")
    set_run_font(identity.add_run("{!Packet.SupplierName}"), name="Poppins SemiBold", size=12.5, bold=True, color=ACCENT)

    title = add_text(
        doc,
        "SUPPLIER CREDIT ACKNOWLEDGEMENT",
        size=21,
        bold=True,
        before=17,
        after=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    shade_paragraph(title, "FFFFFF", border="A6A6A6")

    add_section_label(doc, "Exception summary")
    add_text(
        doc,
        "ClearPacket matched purchase order {!Packet.PurchaseOrderNumber}, invoice {!Packet.InvoiceNumber}, "
        "and the delivery receipt. The invoice lists {!Packet.InvoiceQuantity} units; both the order and receipt "
        "show {!Packet.ReceivedQuantity}.",
        after=5,
    )

    add_section_label(doc, "Verified adjustment")
    decision = add_text(
        doc,
        "Reviewer {!Packet.ReviewerName} approved payment for {!Packet.PayableQuantity} received units at "
        "${!$format(toDecimal(Packet.UnitPrice), 'number', '#,###.00')} each. The supplier credit is "
        "${!$format(toDecimal(Packet.AdjustmentAmount), 'number', '#,###.00')}.",
        after=6,
    )
    shade_paragraph(decision, PALE)

    add_section_label(doc, "Adjustment detail")
    add_expression(
        doc,
        '<mdoc:repeater name="adjustments" value="{!Packet.Adjustments}" variable="item" mode="standard">',
        accent=True,
    )
    add_expression(doc, "{!#item#.ItemSku}  |  {!#item#.Description}")
    add_expression(
        doc,
        "Billed {!#item#.InvoiceQuantity}  |  Received {!#item#.ReceivedQuantity}  |  "
        "Credit ${!$format(toDecimal(#item#.AdjustmentAmount), 'number', '#,###.00')}",
    )
    add_expression(doc, '</mdoc:repeater name="adjustments">', accent=True, after=6)

    total = add_text(
        doc,
        "TOTAL CREDIT: ${!$format(sum(Packet.Adjustments, \"AdjustmentAmount\"), 'number', '#,###.00')}",
        size=11.5,
        bold=True,
        before=2,
        after=7,
    )
    total.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_expression(
        doc,
        '<mdoc:paragraph name="credit-required" hidden="{!$toDecimal(Packet.AdjustmentAmount) <= 0}">',
        accent=True,
    )
    add_text(
        doc,
        "Apex Industrial Supply acknowledges that the invoice must be reduced by "
        "${!$format(toDecimal(Packet.AdjustmentAmount), 'number', '#,###.00')} before payment.",
        after=4,
    )
    add_expression(doc, '</mdoc:paragraph name="credit-required">', accent=True, after=6)

    add_section_label(doc, "Evidence and control")
    add_text(
        doc,
        "Packet {!Packet.PacketId} keeps the three source filenames, extracted values, deterministic checks, "
        "and the human decision dated {!Packet.DecisionDate}. This acknowledgement records the resolution; it "
        "does not authorize payment without the stated review.",
        after=7,
    )

    signoff = doc.add_paragraph()
    set_paragraph_spacing(signoff, before=11, after=0)
    signoff.paragraph_format.tab_stops.add_tab_stop(Inches(3.35))
    set_run_font(signoff.add_run("AP REVIEWER"), size=8.5, bold=True, color=MUTED)
    signoff.add_run("\t")
    set_run_font(signoff.add_run("SUPPLIER ACKNOWLEDGEMENT"), size=8.5, bold=True, color=MUTED)

    lines = doc.add_paragraph()
    set_paragraph_spacing(lines, before=17, after=2)
    lines.paragraph_format.tab_stops.add_tab_stop(Inches(3.35))
    set_run_font(lines.add_run("________________________________"), size=9, color=MUTED)
    lines.add_run("\t")
    set_run_font(lines.add_run("________________________________"), size=9, color=MUTED)

    names = doc.add_paragraph()
    set_paragraph_spacing(names, after=0)
    names.paragraph_format.tab_stops.add_tab_stop(Inches(3.35))
    set_run_font(names.add_run("{!Packet.ReviewerName}"), size=9.5)
    names.add_run("\t")
    set_run_font(names.add_run("{!Packet.SupplierName}"), size=9.5)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        footer_paragraph.add_run("ClearPacket evidence  /  Doctavian document generation"),
        size=7.5,
        color=MUTED,
    )

    doc.core_properties.title = "ClearPacket supplier credit acknowledgement"
    doc.core_properties.subject = "Doctavian generation template"
    doc.core_properties.author = "ClearPacket"
    doc.core_properties.keywords = "ClearPacket, Doctavian, invoice, credit acknowledgement"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("reference", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.reference, args.output)


if __name__ == "__main__":
    main()
